import docx
from docx.shared import Inches


doc = docx.Document()
doc.add_paragraph('Hello world!')
doc.add_picture('./test.jpg' , width=Inches(1.5748), height=Inches(2.3622))
doc.add_paragraph('Hello world! x2')
doc.save('demo3.docx')