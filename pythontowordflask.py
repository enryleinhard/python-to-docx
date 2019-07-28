from flask import Flask, request
from flask_restful import Resource, Api
from flask_cors import CORS, cross_origin
import docx
from docx.shared import Inches

application = app = Flask(__name__)
api = Api(app)

class HelloWorld(Resource):
    def get(self):
        return "Hello World"

api.add_resource(HelloWorld, '/')

class text(Resource):
    def get(self, text, text2):

        doc = docx.Document('pythontowordflask.docx')
        doc.add_paragraph(text)
        doc.add_paragraph(text2)
        doc.save('pythontowordflask.docx') 

api.add_resource(text, '/text/<text>/<text2>')

class getImage(Resource):
    def get(self, gambar):
        doc = docx.Document()
        doc.add_picture('./'+gambar , width=Inches(1.5748), height=Inches(2.3622))
        doc.save('pythontowordflask.docx')

api.add_resource(getImage, '/image/<gambar>')

# harus masukin image dl baru masukin text.

if __name__ == '__main__':
    app.run(debug=True)